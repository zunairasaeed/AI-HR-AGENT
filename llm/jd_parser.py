############################working fine below code handles pararagh jd
# import os
# import json
# from docx import Document
# import fitz  # PyMuPDF
# from tqdm import tqdm
# from app.llm_utils import call_gpt

# # 📁 Folder paths
# RAW_JD_FOLDER = os.path.join("data", "jobs", "raw_jd")
# STRUCTURED_JD_FOLDER = os.path.join("data", "jobs", "parsed_jd")  # ✅ Corrected output path

# os.makedirs(STRUCTURED_JD_FOLDER, exist_ok=True)
# os.makedirs(RAW_JD_FOLDER, exist_ok=True)

# # -------------------------------
# # 📄 Read text from file
# # -------------------------------
# def extract_text_from_file(filepath):
#     if filepath.endswith(".pdf"):
#         text = ""
#         with fitz.open(filepath) as doc:
#             for page in doc:
#                 text += page.get_text()
#         return text
#     elif filepath.endswith(".docx"):
#         doc = Document(filepath)
#         return "\n".join([p.text for p in doc.paragraphs])
#     else:
#         return ""  # unsupported

# # -------------------------------
# # 🤖 LLM: Structure the JD
# # -------------------------------
# def extract_jd_data_llm(text):
#     system_prompt = "You are a helpful assistant that structures unformatted job descriptions."

#     # 🔒 Optional: Truncate very long JDs
#     if len(text) > 8000:
#         text = text[:8000]

#     user_prompt = f"""
# Parse the following job description and return structured data as valid JSON.

# Return the following keys:

# {{
#   "job_title": "",
#   "department": "",
#   "location": "",
#   "employment_type": "",
#   "must_have_skills": [],
#   "nice_to_have_skills": [],
#   "tools": [],
#   "responsibilities": [],
#   "requirements": [],
#   "experience_level": "",
#   "education_required": "",
#   "keywords": []
# }}

# Ensure no duplicate or redundant values and output only valid JSON (no extra text).

# Job Description:
# \"\"\"{text}\"\"\"
#     """

#     response = call_gpt(system_prompt, user_prompt)

#     # 🧠 Log raw GPT response
#     print("\n🧠 Raw GPT Response:\n", response)

#     # ✅ Strip markdown-style code blocks before parsing
#     try:
#         if not response.strip():
#             raise ValueError("Empty response from LLM.")

#         # Remove ```json or ``` from start/end
#         if response.strip().startswith("```"):
#             response = "\n".join(
#                 line for line in response.strip().splitlines()
#                 if not line.strip().startswith("```")
#             )

#         return json.loads(response)

#     except Exception as e:
#         print("❌ JSON parse error:", e)
#         return {}

# # -------------------------------
# # 🔁 Process each JD
# # -------------------------------

# def run_job_parsing():
#     for filename in tqdm(os.listdir(RAW_JD_FOLDER)):
#         filepath = os.path.join(RAW_JD_FOLDER, filename)
#         if not (filename.endswith(".pdf") or filename.endswith(".docx")):
#             continue

#         text = extract_text_from_file(filepath)
#         print(f"\n📄 Preview from {filename}:\n{text[:1000]}")

#         structured_data = extract_jd_data_llm(text)

#         if structured_data:
#             safe_name = os.path.splitext(filename)[0].replace(" ", "_")
#             output_path = os.path.join(STRUCTURED_JD_FOLDER, f"{safe_name}.json")
#             with open(output_path, "w", encoding="utf-8") as f_out:
#                 json.dump(structured_data, f_out, indent=2, ensure_ascii=False)
#             print(f"✅ Parsed: {filename} → {output_path}")
#         else:
#             print(f"⚠️ Skipped saving empty output for {filename}")

#     print("\n🎉 All job descriptions processed and structured.")

# # 🔘 Optional CLI execution
# if __name__ == "__main__":
#     run_job_parsing()











#####################code contains combined folder of all te jods descri[ption in 1 file ]

import os
import json
from docx import Document
import fitz  # PyMuPDF
from tqdm import tqdm
from app.llm_utils import call_gpt

# 📁 Folder paths
RAW_JD_FOLDER = os.getenv("RAW_JD_DIR", os.path.join("data", "jobs", "raw_jd"))
STRUCTURED_JD_FOLDER = os.getenv("PARSED_JD_DIR", os.path.join("data", "jobs", "parsed_jd"))
RAW_PARSED_JD_FOLDER = os.path.join(os.path.dirname(STRUCTURED_JD_FOLDER), "raw_parsed_jd")
COMBINED_JD_FOLDER = os.getenv("COMBINED_JD_DIR", os.path.join(STRUCTURED_JD_FOLDER, "combined"))
SESSION_LOCK_PATH = os.getenv('SESSION_LOCK_PATH', os.path.join('data', 'session_lock.json'))

os.makedirs(STRUCTURED_JD_FOLDER, exist_ok=True)
os.makedirs(RAW_JD_FOLDER, exist_ok=True)
os.makedirs(RAW_PARSED_JD_FOLDER, exist_ok=True)
os.makedirs(COMBINED_JD_FOLDER, exist_ok=True)

# -------------------------------
# 📄 Read text from file
# -------------------------------
def extract_text_from_file(filepath):
    if filepath.endswith(".pdf"):
        text = ""
        with fitz.open(filepath) as doc:
            for page in doc:
                text += page.get_text()
        return text
    elif filepath.endswith(".docx"):
        doc = Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return ""  # unsupported

# -------------------------------
# 🤖 LLM: Structure JD normally
# -------------------------------
def extract_jd_data_llm(text):
    system_prompt = "You are a helpful assistant that structures unformatted job descriptions."
    if len(text) > 8000:
        text = text[:8000]

    user_prompt = f"""
Parse the following job description and return structured data as valid JSON.

Required schema:
{{
  "job_title": "",
  "department": "",
  "location": "",
  "employment_type": "",
  "must_have_skills": [],
  "nice_to_have_skills": [],
  "tools": [],
  "responsibilities": [],
  "requirements": [],
  "experience_level": "",
  "education_required": "",
  "keywords": []
}}

Ensure no duplicate values and output only valid JSON.

Job Description:
\"\"\"{text}\"\"\""""
    response = call_gpt(system_prompt, user_prompt)

    try:
        if not response.strip():
            raise ValueError("Empty response from LLM.")
        if response.strip().startswith("```"):
            response = "\n".join(line for line in response.splitlines() if not line.strip().startswith("```"))
        return json.loads(response)
    except Exception as e:
        print("❌ JSON parse error:", e)
        return {}

# -------------------------------
# 🤖 NEW: Paragraph-style JD parsing
# -------------------------------
def extract_paragraph_jd(text, filename):
    system_prompt = "You are an expert HR assistant specializing in structuring unformatted job descriptions."
    if len(text) > 8000:
        text = text[:8000]

    user_prompt = f"""
The following job description is written in a paragraph format.

Step 1: Extract all possible raw information (skills, requirements, responsibilities, tools, qualifications) clearly.

Step 2: Transform it into the structured JSON format with the schema below:

{{
  "job_title": "",
  "department": "",
  "location": "",
  "employment_type": "",
  "must_have_skills": [],
  "nice_to_have_skills": [],
  "tools": [],
  "responsibilities": [],
  "requirements": [],
  "experience_level": "",
  "education_required": "",
  "keywords": []
}}

Return ONLY valid JSON.

JD Text:
\"\"\"{text}\"\"\""""
    response = call_gpt(system_prompt, user_prompt)

    try:
        if not response.strip():
            raise ValueError("Empty response from LLM.")
        if response.strip().startswith("```"):
            response = "\n".join(line for line in response.splitlines() if not line.strip().startswith("```"))

        parsed_json = json.loads(response)

        # Save raw parsed JD separately
        raw_output_path = os.path.join(RAW_PARSED_JD_FOLDER, f"{os.path.splitext(filename)[0]}_raw.json")
        with open(raw_output_path, "w", encoding="utf-8") as f_raw:
            json.dump(parsed_json, f_raw, indent=2, ensure_ascii=False)
        print(f"📝 Raw paragraph JD saved: {raw_output_path}")

        return parsed_json

    except Exception as e:
        print("❌ JSON parse error (paragraph JD):", e)
        return {}

# -------------------------------
# 🔁 Process each JD
# -------------------------------
def _load_session_uploaded_jds():
    try:
        if os.path.exists(SESSION_LOCK_PATH):
            with open(SESSION_LOCK_PATH, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
            return set(data.get("uploaded_jds", []))
    except Exception:
        pass
    return None

def _update_session_parsed_jds(parsed_filenames):
    try:
        os.makedirs(os.path.dirname(SESSION_LOCK_PATH), exist_ok=True)
        data = {}
        if os.path.exists(SESSION_LOCK_PATH):
            with open(SESSION_LOCK_PATH, "r", encoding="utf-8") as f:
                try:
                    data = json.load(f) or {}
                except json.JSONDecodeError:
                    data = {}
        data["parsed_jds"] = sorted(list(set(parsed_filenames)))
        with open(SESSION_LOCK_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"⚠️ Could not update session lock (parsed_jds): {e}")

def run_job_parsing():
    filter_set = _load_session_uploaded_jds()
    parsed_out_files = []
    for filename in tqdm(os.listdir(RAW_JD_FOLDER)):
        filepath = os.path.join(RAW_JD_FOLDER, filename)
        if not (filename.endswith(".pdf") or filename.endswith(".docx")):
            continue
        if filter_set is not None and filename not in filter_set:
            continue

        text = extract_text_from_file(filepath)
        print(f"\n📄 Preview from {filename}:\n{text[:500]}")

        # ✅ Decide: Paragraph JD or Normal JD
        if len(text.split("\n")) <= 2:  
            structured_data = extract_paragraph_jd(text, filename)
        else:
            structured_data = extract_jd_data_llm(text)

        if structured_data:
            safe_name = os.path.splitext(filename)[0].replace(" ", "_")
            output_path = os.path.join(STRUCTURED_JD_FOLDER, f"{safe_name}.json")
            with open(output_path, "w", encoding="utf-8") as f_out:
                json.dump(structured_data, f_out, indent=2, ensure_ascii=False)
            print(f"✅ Parsed: {filename} → {output_path}")
            parsed_out_files.append(os.path.basename(output_path))
        else:
            print(f"⚠️ Skipped saving empty output for {filename}")

    print("\n🎉 All job descriptions processed and structured.")

    # ✅ After parsing → Combine all parsed JDs
    _update_session_parsed_jds(parsed_out_files)
    combine_all_parsed_jds()

# -------------------------------
# 🆕 Combine all structured JDs
# -------------------------------
def combine_all_parsed_jds():
    all_jobs = []

    session_files = None
    try:
        if os.path.exists(SESSION_LOCK_PATH):
            with open(SESSION_LOCK_PATH, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
            session_files = set(data.get("parsed_jds", []))
    except Exception:
        session_files = None

    for file in os.listdir(STRUCTURED_JD_FOLDER):
        if not file.endswith(".json") or file == "combined":
            continue
        if session_files is not None and file not in session_files:
            continue
        filepath = os.path.join(STRUCTURED_JD_FOLDER, file)

        with open(filepath, "r", encoding="utf-8") as f:
            jd_content = json.load(f)

        if isinstance(jd_content, dict) and "jobs" in jd_content:
            all_jobs.extend(jd_content["jobs"])
        elif isinstance(jd_content, list):
            all_jobs.extend(jd_content)
        elif isinstance(jd_content, dict) and "job_title" in jd_content:
            all_jobs.append(jd_content)
        elif isinstance(jd_content, dict):
            for jt, job in jd_content.items():
                job["job_title"] = jt
                all_jobs.append(job)

    # Save into one master file
    combined_path = os.path.join(COMBINED_JD_FOLDER, "all_jobs.json")
    with open(combined_path, "w", encoding="utf-8") as f_out:
        json.dump({"jobs": all_jobs}, f_out, indent=2, ensure_ascii=False)

    print(f"\n📦 Combined {len(all_jobs)} jobs → {combined_path}")

# 🔘 CLI
if __name__ == "__main__":
    run_job_parsing()
